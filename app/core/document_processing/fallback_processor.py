#!/usr/bin/env python3
"""
Fallback Document Processor
Handles unsupported formats and provides graceful degradation
"""

import logging
import asyncio
from typing import Dict, Any, List, Optional, Union, BinaryIO
from datetime import datetime
import tempfile
import os

from app.core.document_processing.base import (
    DocumentProcessor, ProcessingResult, DocumentType,
    DocumentMetadata, StructuredContent, TableData, ImageData
)
from app.core.document_processing.exceptions import (
    ProcessingError, UnsupportedFormatError
)

logger = logging.getLogger(__name__)


class FallbackProcessor(DocumentProcessor):
    """Fallback processor for unsupported document formats."""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize fallback processor."""
        super().__init__(config)
        self.max_file_size = config.get("max_file_size", 50 * 1024 * 1024)  # 50MB
        self.timeout_seconds = config.get("timeout_seconds", 60)
    
    def supports_format(self, file_type: DocumentType) -> bool:
        """Fallback processor supports all formats as last resort."""
        return True
    
    def get_supported_formats(self) -> List[DocumentType]:
        """Get list of all document formats (fallback for all)."""
        return list(DocumentType)
    
    async def process_document(
        self, 
        file_data: Union[bytes, BinaryIO], 
        file_type: DocumentType,
        filename: Optional[str] = None,
        **kwargs
    ) -> ProcessingResult:
        """
        Process document using fallback methods.
        
        Args:
            file_data: Document data as bytes or file-like object
            file_type: Type of document to process
            filename: Original filename (optional)
            **kwargs: Additional processing options
            
        Returns:
            ProcessingResult with basic extracted content
        """
        start_time = datetime.now()
        
        try:
            # Validate input
            self.validate_input(file_data, file_type, self.max_file_size)
            
            # Convert file data to bytes if needed
            if isinstance(file_data, bytes):
                data_bytes = file_data
            else:
                data_bytes = file_data.read()
            
            logger.info(f"Processing {filename} with fallback processor for {file_type.value}")
            
            # Try format-specific fallback processing
            if file_type == DocumentType.PDF:
                result = await self._process_pdf_fallback(data_bytes, filename)
            elif file_type == DocumentType.DOCX:
                result = await self._process_docx_fallback(data_bytes, filename)
            elif file_type == DocumentType.HTML:
                result = await self._process_html_fallback(data_bytes, filename)
            elif file_type == DocumentType.TXT:
                result = await self._process_text_fallback(data_bytes, filename)
            elif file_type == DocumentType.RTF:
                result = await self._process_rtf_fallback(data_bytes, filename)
            else:
                # Generic text extraction attempt
                result = await self._process_generic_fallback(data_bytes, filename, file_type)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            result.processing_time = processing_time
            
            # Add fallback processing metadata
            if result.raw_data is None:
                result.raw_data = {}
            
            result.raw_data.update({
                "processor_name": "fallback",
                "fallback_method": f"{file_type.value}_fallback",
                "processing_warnings": ["Using fallback processing - limited functionality"]
            })
            
            logger.info(f"Fallback processing completed for {filename} in {processing_time:.2f}s")
            return result
            
        except Exception as e:
            logger.error(f"Fallback processing failed for {filename}: {e}")
            
            # Return minimal result even on failure
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessingResult(
                success=False,
                content=StructuredContent(text=""),
                metadata=DocumentMetadata(
                    title=filename or "Unknown Document",
                    custom_properties={
                        "processing_error": str(e),
                        "fallback_attempted": True
                    }
                ),
                processing_time=processing_time,
                errors=[f"Fallback processing failed: {str(e)}"],
                raw_data={
                    "processor_name": "fallback",
                    "error_details": str(e)
                }
            )
    
    async def _process_pdf_fallback(self, data_bytes: bytes, filename: Optional[str]) -> ProcessingResult:
        """Fallback PDF processing using PyPDF2 or similar."""
        try:
            # Try PyPDF2 first
            try:
                import PyPDF2
                from io import BytesIO
                
                pdf_reader = PyPDF2.PdfReader(BytesIO(data_bytes))
                text_content = ""
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                metadata = DocumentMetadata(
                    title=filename or "PDF Document",
                    page_count=len(pdf_reader.pages),
                    word_count=len(text_content.split()),
                    custom_properties={
                        "extraction_method": "PyPDF2",
                        "fallback_processor": True
                    }
                )
                
                return ProcessingResult(
                    success=True,
                    content=StructuredContent(text=text_content.strip()),
                    metadata=metadata
                )
                
            except ImportError:
                logger.warning("PyPDF2 not available for PDF fallback")
                
            # Try pdfplumber as alternative
            try:
                import pdfplumber
                from io import BytesIO
                
                text_content = ""
                page_count = 0
                
                with pdfplumber.open(BytesIO(data_bytes)) as pdf:
                    for page in pdf.pages:
                        text_content += page.extract_text() or ""
                        text_content += "\n"
                        page_count += 1
                
                metadata = DocumentMetadata(
                    title=filename or "PDF Document",
                    page_count=page_count,
                    word_count=len(text_content.split()),
                    custom_properties={
                        "extraction_method": "pdfplumber",
                        "fallback_processor": True
                    }
                )
                
                return ProcessingResult(
                    success=True,
                    content=StructuredContent(text=text_content.strip()),
                    metadata=metadata
                )
                
            except ImportError:
                logger.warning("pdfplumber not available for PDF fallback")
            
            # If no PDF libraries available, return basic result
            return self._create_basic_result(data_bytes, filename, "PDF (no extraction libraries)")
            
        except Exception as e:
            logger.error(f"PDF fallback processing failed: {e}")
            return self._create_basic_result(data_bytes, filename, f"PDF (error: {str(e)})")
    
    async def _process_docx_fallback(self, data_bytes: bytes, filename: Optional[str]) -> ProcessingResult:
        """Fallback DOCX processing using python-docx."""
        try:
            import docx
            from io import BytesIO
            
            doc = docx.Document(BytesIO(data_bytes))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            metadata = DocumentMetadata(
                title=filename or "DOCX Document",
                word_count=len(text_content.split()),
                custom_properties={
                    "extraction_method": "python-docx",
                    "fallback_processor": True,
                    "paragraph_count": len(doc.paragraphs)
                }
            )
            
            return ProcessingResult(
                success=True,
                content=StructuredContent(text=text_content.strip()),
                metadata=metadata
            )
            
        except ImportError:
            logger.warning("python-docx not available for DOCX fallback")
            return self._create_basic_result(data_bytes, filename, "DOCX (no extraction library)")
        except Exception as e:
            logger.error(f"DOCX fallback processing failed: {e}")
            return self._create_basic_result(data_bytes, filename, f"DOCX (error: {str(e)})")
    
    async def _process_html_fallback(self, data_bytes: bytes, filename: Optional[str]) -> ProcessingResult:
        """Fallback HTML processing using BeautifulSoup."""
        try:
            from bs4 import BeautifulSoup
            
            html_content = data_bytes.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text(separator='\n', strip=True)
            
            # Extract title
            title_tag = soup.find('title')
            title = title_tag.get_text() if title_tag else (filename or "HTML Document")
            
            metadata = DocumentMetadata(
                title=title,
                word_count=len(text_content.split()),
                custom_properties={
                    "extraction_method": "BeautifulSoup",
                    "fallback_processor": True,
                    "html_tags_found": len(soup.find_all())
                }
            )
            
            return ProcessingResult(
                success=True,
                content=StructuredContent(text=text_content),
                metadata=metadata
            )
            
        except ImportError:
            logger.warning("BeautifulSoup not available for HTML fallback")
            # Try basic HTML tag removal
            try:
                import re
                html_content = data_bytes.decode('utf-8', errors='ignore')
                text_content = re.sub(r'<[^>]+>', '', html_content)
                text_content = re.sub(r'\s+', ' ', text_content).strip()
                
                return ProcessingResult(
                    success=True,
                    content=StructuredContent(text=text_content),
                    metadata=DocumentMetadata(
                        title=filename or "HTML Document",
                        word_count=len(text_content.split()),
                        custom_properties={
                            "extraction_method": "regex_html_strip",
                            "fallback_processor": True
                        }
                    )
                )
            except Exception:
                pass
            
            return self._create_basic_result(data_bytes, filename, "HTML (no parsing library)")
        except Exception as e:
            logger.error(f"HTML fallback processing failed: {e}")
            return self._create_basic_result(data_bytes, filename, f"HTML (error: {str(e)})")
    
    async def _process_text_fallback(self, data_bytes: bytes, filename: Optional[str]) -> ProcessingResult:
        """Fallback text processing."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
            text_content = None
            
            for encoding in encodings:
                try:
                    text_content = data_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                text_content = data_bytes.decode('utf-8', errors='ignore')
            
            metadata = DocumentMetadata(
                title=filename or "Text Document",
                word_count=len(text_content.split()),
                character_count=len(text_content),
                custom_properties={
                    "extraction_method": "text_decode",
                    "fallback_processor": True
                }
            )
            
            return ProcessingResult(
                success=True,
                content=StructuredContent(text=text_content),
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Text fallback processing failed: {e}")
            return self._create_basic_result(data_bytes, filename, f"Text (error: {str(e)})")
    
    async def _process_rtf_fallback(self, data_bytes: bytes, filename: Optional[str]) -> ProcessingResult:
        """Fallback RTF processing."""
        try:
            # Try striprtf library
            try:
                from striprtf.striprtf import rtf_to_text
                rtf_content = data_bytes.decode('utf-8', errors='ignore')
                text_content = rtf_to_text(rtf_content)
                
                metadata = DocumentMetadata(
                    title=filename or "RTF Document",
                    word_count=len(text_content.split()),
                    custom_properties={
                        "extraction_method": "striprtf",
                        "fallback_processor": True
                    }
                )
                
                return ProcessingResult(
                    success=True,
                    content=StructuredContent(text=text_content),
                    metadata=metadata
                )
                
            except ImportError:
                logger.warning("striprtf not available for RTF fallback")
                # Basic RTF text extraction
                rtf_content = data_bytes.decode('utf-8', errors='ignore')
                # Remove RTF control codes (basic approach)
                import re
                text_content = re.sub(r'\\[a-z]+\d*\s?', '', rtf_content)
                text_content = re.sub(r'[{}]', '', text_content)
                text_content = re.sub(r'\s+', ' ', text_content).strip()
                
                return ProcessingResult(
                    success=True,
                    content=StructuredContent(text=text_content),
                    metadata=DocumentMetadata(
                        title=filename or "RTF Document",
                        word_count=len(text_content.split()),
                        custom_properties={
                            "extraction_method": "basic_rtf_strip",
                            "fallback_processor": True
                        }
                    )
                )
                
        except Exception as e:
            logger.error(f"RTF fallback processing failed: {e}")
            return self._create_basic_result(data_bytes, filename, f"RTF (error: {str(e)})")
    
    async def _process_generic_fallback(
        self, 
        data_bytes: bytes, 
        filename: Optional[str], 
        file_type: DocumentType
    ) -> ProcessingResult:
        """Generic fallback processing for unknown formats."""
        try:
            # Try to extract any readable text
            text_content = data_bytes.decode('utf-8', errors='ignore')
            
            # Clean up binary content
            import re
            text_content = re.sub(r'[^\x20-\x7E\n\r\t]', '', text_content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            metadata = DocumentMetadata(
                title=filename or f"Unknown {file_type.value} Document",
                word_count=len(text_content.split()) if text_content else 0,
                custom_properties={
                    "extraction_method": "generic_text_extraction",
                    "fallback_processor": True,
                    "original_format": file_type.value
                }
            )
            
            return ProcessingResult(
                success=bool(text_content),
                content=StructuredContent(text=text_content),
                metadata=metadata,
                warnings=["Generic text extraction used - content may be incomplete"]
            )
            
        except Exception as e:
            logger.error(f"Generic fallback processing failed: {e}")
            return self._create_basic_result(data_bytes, filename, f"Generic (error: {str(e)})")
    
    def _create_basic_result(
        self, 
        data_bytes: bytes, 
        filename: Optional[str], 
        method: str
    ) -> ProcessingResult:
        """Create a basic processing result when extraction fails."""
        return ProcessingResult(
            success=False,
            content=StructuredContent(text=""),
            metadata=DocumentMetadata(
                title=filename or "Unknown Document",
                file_size=len(data_bytes),
                custom_properties={
                    "extraction_method": method,
                    "fallback_processor": True,
                    "extraction_failed": True
                }
            ),
            warnings=[f"Content extraction failed using {method}"],
            errors=["Unable to extract readable content from document"]
        )
